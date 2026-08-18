import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

md_file = r"d:\University\Summer 2026\DN Smart Trade ERP AI Platform\Project Report\Arunavo_Full_Report_Sections.md"
output_file = r"d:\University\Summer 2026\DN Smart Trade ERP AI Platform\Project Report\Arunavo_Project_Report.docx"
image_dir = r"d:\University\Summer 2026\DN Smart Trade ERP AI Platform\Project Report"

diagram_descriptions = {
    "Architecture_Diagram_drawio.png": "Description: This diagram illustrates the high-level layered architecture of the DN Smart Trade ERP platform, showcasing the separation between the presentation layer (Next.js), application layer (Server Actions), data layer (Supabase), and AI layer (Groq API).",
    "SequenceLoginAuthentication1_drawio.png": "Description: This sequence diagram details the login and authentication flow, demonstrating how the system validates user credentials via Supabase Auth and performs role-based redirection.",
    "Shipment_Creation_drawio.png": "Description: This sequence diagram outlines the shipment creation process, showing the interactions between the user, the Next.js Server Actions, and the database for saving shipment records and timelines.",
    "Bill_of_Entry__BOE__drawio.png": "Description: This sequence diagram depicts the creation of a Bill of Entry, highlighting the data validation steps, duty calculation, and database insertion processes.",
    "BOE_Status_Transition__with_role-based_override__drawio.png": "Description: This diagram maps out the state machine transitions for a Bill of Entry, emphasizing the role-based access control where Admins can override the standard lifecycle while Employees cannot.",
    "AI-Assisted_HS_Code_Classification_drawio.png": "Description: This sequence diagram explains the AI-assisted HS Code classification feature, detailing how the user's query is processed by the Groq LLM to return ranked, confidence-scored candidate matches.",
    "document-upload-PLACEHOLDER.png": "Description: This sequence diagram shows the document upload workflow, including file storage via Supabase Storage and metadata insertion with constraints to link documents to specific entities.",
    "Role-Based_Notification_Dispatch_drawio.png": "Description: This diagram illustrates the role-based notification dispatch system, describing how alerts are routed and persisted to specific users based on roles and system events."
}

image_mapping = {
    "Architecture_Diagram_drawio.png": "Architecture Diagram.drawio.png",
    "SequenceLoginAuthentication1_drawio.png": "SequenceLoginAuthentication1.drawio.png",
    "Shipment_Creation_drawio.png": "Shipment Creation.drawio.png",
    "Bill_of_Entry__BOE__drawio.png": "Bill of Entry (BOE).drawio.png",
    "BOE_Status_Transition__with_role-based_override__drawio.png": "BOE Status Transition (with role-based override).drawio.png",
    "AI-Assisted_HS_Code_Classification_drawio.png": "AI-Assisted HS Code Classification.drawio (1).png",
    "document-upload-PLACEHOLDER.png": "Document Upload.drawio.png",
    "Role-Based_Notification_Dispatch_drawio.png": "Role-Based Notification Dispatch.drawio.png"
}

doc = Document()

# Configure Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
style.paragraph_format.space_after = Pt(10)
style.paragraph_format.line_spacing = 1.15

# Headings
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(12)
h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)
h2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)
h3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# List bullets
bullet_style = doc.styles['List Bullet']
bullet_style.font.name = 'Times New Roman'
bullet_style.font.size = Pt(12)
bullet_style.paragraph_format.line_spacing = 1.15
bullet_style.paragraph_format.space_after = Pt(6)
bullet_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def apply_table_style(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

with open(md_file, "r", encoding="utf-8") as f:
    lines = f.readlines()

def parse_bold_italic(p, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)

in_table = False
table = None
in_comment = False

for line in lines:
    original_line = line
    line = line.strip()
    
    if line.startswith('<!--'):
        in_comment = True
    if in_comment:
        if '-->' in original_line:
            in_comment = False
        continue

    if not line:
        in_table = False
        continue

    if line.startswith('|'):
        if not in_table:
            in_table = True
            cols = [c.strip() for c in line.split('|')[1:-1]]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(cols):
                hdr_cells[i].text = col_name
        else:
            if '---' in line:
                continue 
            cols = [c.strip() for c in line.split('|')[1:-1]]
            if len(cols) == len(table.columns):
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cols):
                    row_cells[i].text = cell_text
        # Re-apply styles after table row addition to maintain formatting
        apply_table_style(table)
        continue
    else:
        in_table = False

    if line.startswith('#'):
        # handle heading
        level_str = line.split(' ')[0]
        if all(c == '#' for c in level_str):
            level = len(level_str)
            text = line[level:].strip()
            # If level is deeper than 3, fallback to 3
            if level > 3:
                level = 3
            p = doc.add_paragraph(style=f'Heading {level}')
            parse_bold_italic(p, text)
        else:
            p = doc.add_paragraph()
            parse_bold_italic(p, line)
    elif line.startswith('- '):
        text = line[2:]
        p = doc.add_paragraph(style='List Bullet')
        parse_bold_italic(p, text)
    elif line.startswith('!['):
        match = re.search(r'\!\[.*?\]\((.*?)\)', line)
        if match:
            img_md_name = match.group(1)
            actual_name = image_mapping.get(img_md_name, img_md_name)
            img_path = os.path.join(image_dir, actual_name)
            if os.path.exists(img_path):
                # Ensure we have space before the image
                doc.add_paragraph()
                p_img = doc.add_paragraph()
                p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Inches(6.0))
                
                desc = diagram_descriptions.get(img_md_name, "Description: Diagram showing system workflow.")
                p_desc = doc.add_paragraph()
                p_desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p_desc.paragraph_format.space_before = Pt(6)
                p_desc.paragraph_format.space_after = Pt(24)
                run_desc = p_desc.add_run(desc)
                run_desc.italic = True
                run_desc.font.name = 'Times New Roman'
                run_desc.font.size = Pt(11) # Description slightly smaller
            else:
                doc.add_paragraph(f"[Image not found: {actual_name}]")
    else:
        p = doc.add_paragraph()
        parse_bold_italic(p, line)

doc.save(output_file)
print(f"Successfully generated formatted {output_file}")
